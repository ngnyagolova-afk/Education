import sys, random
sys.stdout.reconfigure(encoding='utf-8')
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

QUESTIONS = [
    ('Раздел 6', 'Какво представлява файловата система?',
     ['Физическо устройство за съхранение на данни',
      'Механизъм, контролиращ организацията, достъпа и съхранението на данни',
      'Програма за управление на процеси',
      'Протокол за мрежова комуникация'], 1),
    ('Раздел 6', 'На колко байта е стандартният размер на един сектор на твърдия диск?',
     ['256 байта', '4096 байта', '1024 байта', '512 байта'], 3),
    ('Раздел 6', 'Файловата система ext4 е характерна предимно за:',
     ['Windows', 'macOS', 'Linux', 'FreeBSD'], 2),
    ('Раздел 6', 'Каква е максималната дискова памет, поддържана от MBR?',
     ['512 GB', '4 TB', '8 ZB', '2 TB'], 3),
    ('Раздел 6', 'Колко първични дяла поддържа MBR схемата?',
     ['2', '8', '128', '4'], 3),
    ('Раздел 6', 'GPT поддържа до колко дяла?',
     ['4', '32', '128', '64'], 2),
    ('Раздел 6', 'Коя команда в Linux монтира файлова система?',
     ['lsblk', 'fsck', 'fdisk', 'mount'], 3),
    ('Раздел 6', 'Коя команда демонтира файлова система в Linux?',
     ['dismount', 'detach', 'unmount', 'umount'], 3),
    ('Раздел 6', 'Коя команда показва списъка с блокови устройства в Linux?',
     ['ls -l /dev', 'blkid', 'fdisk -l', 'lsblk'], 3),
    ('Раздел 6', 'Каква е функцията на командата fsck?',
     ['Форматира файловата система',
      'Монтира файловата система',
      'Показва размера на файловата система',
      'Проверява и поправя консистентността на файловата система'], 3),
    ('Раздел 6', 'Коя файлова система е разработена от Microsoft?',
     ['ext4', 'ReiserFS', 'Btrfs', 'NTFS'], 3),
    ('Раздел 6', 'Коя команда показва UUID и типа на файловата система на даден дял?',
     ['lsblk', 'mount -l', 'fdisk -l', 'blkid'], 3),
    ('Раздел 6', 'FAT32 файловата система се използва предимно за:',
     ['Сървърни Linux системи', 'Основни Windows системни дискове',
      'Мрежови файлови системи', 'Flash памети и преносими носители'], 3),
    ('Раздел 6', 'Какво е MBR (Master Boot Record)?',
     ['Модерна таблица на дялове с поддръжка за 128 дяла',
      'Вид файлова система за Linux',
      'Протокол за форматиране на диск',
      'Първият сектор на диска, използван за начално зареждане'], 3),
    ('Раздел 6', 'Коя команда се използва за управление на дискови дялове?',
     ['lsblk', 'blkid', 'mount', 'fdisk'], 3),
    ('Раздел 7', 'Какво представлява shebang линията в shell скрипт?',
     ['Коментар, обясняващ целта на скрипта',
      'Последният ред на скрипта',
      'Декларация на глобална променлива',
      'Първият ред, указващ кой интерпретатор да се използва (напр. #!/bin/bash)'], 3),
    ('Раздел 7', 'Коя команда дава право на изпълнение на shell скрипт?',
     ['chown +x script.sh', 'exec script.sh', 'run +x script.sh', 'chmod +x script.sh'], 3),
    ('Раздел 7', 'Как се изпълнява скрипт, когато НЕ е в PATH директорията?',
     ['script.sh', 'run script.sh', 'bash -run script.sh', './script.sh'], 3),
    ('Раздел 7', 'Как правилно се декларира променлива в bash?',
     ['var name = value', '$name=value', 'let name = value', 'name=value'], 3),
    ('Раздел 7', 'Как се получава стойността на променлива name в bash?',
     ['{name}', '@name', '#name', '$name'], 3),
    ('Раздел 7', 'Коя команда чете вход от потребителя и го записва в променлива?',
     ['input', 'get', 'scan', 'read'], 3),
    ('Раздел 7', 'Какво прави операторът >> при пренасочване на изход в bash?',
     ['Изтрива файла и записва новото съдържание',
      'Чете от файл',
      'Пренасочва само съобщенията за грешки',
      'Добавя изхода в края на файл, без да го изтрива'], 3),
    ('Раздел 7', 'С каква ключова дума завършва if блокът в bash?',
     ['end', 'endif', 'done', 'fi'], 3),
    ('Раздел 7', 'Кой оператор сравнява две числа за РАВЕНСТВО в bash?',
     ['==', '=', '-is', '-eq'], 3),
    ('Раздел 7', 'Какво прави командата elif в bash?',
     ['Завършва if блока',
      'Изпълнява код ако всички условия са неверни',
      'Повтаря условната проверка',
      'Проверява допълнително условие, ако предишното е невярно'], 3),
    ('Раздел 7', 'Какво прави цикълът while в bash?',
     ['Изпълнява код точно определен брой пъти',
      'Изпълнява код докато условието е НЕВЯРНО',
      'Изпълнява код само веднъж',
      'Изпълнява код докато условието е ВЯРНО'], 3),
    ('Раздел 7', 'С каква ключова дума завършва for/while цикъл в bash?',
     ['end', 'fi', 'endfor', 'done'], 3),
    ('Раздел 7', 'Какво представлява $# в bash скрипт?',
     ['Стойността на последния аргумент',
      'Пътят до изпълнявания скрипт',
      'Изходният код на последната команда',
      'Броят на аргументите, подадени на скрипта'], 3),
    ('Раздел 7', 'Какво прави командата break в цикъл?',
     ['Прескача текущата итерация и преминава към следващата',
      'Завършва целия скрипт',
      'Извежда съобщение за грешка',
      'Прекратява изпълнението на целия цикъл'], 3),
    ('Раздел 7', 'Какво прави командата continue в цикъл?',
     ['Прекратява изпълнението на цикъла',
      'Завършва скрипта с код 0',
      'Повтаря текущата итерация отначало',
      'Прескача остатъка от текущата итерация и преминава към следващата'], 3),
]

LETTERS = ['А', 'Б', 'В', 'Г']
NUM_VARIANTS = 26

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def build_variant(seed):
    """Връща списък от (section, question, shuffled_opts, correct_letter_index)"""
    rng = random.Random(seed)
    q_indices = list(range(len(QUESTIONS)))
    rng.shuffle(q_indices)
    result = []
    for qi in q_indices:
        sec, q, opts, correct = QUESTIONS[qi]
        correct_text = opts[correct]
        shuffled = opts[:]
        rng.shuffle(shuffled)
        new_correct = shuffled.index(correct_text)
        result.append((sec, q, shuffled, new_correct))
    return result

# ── Генериране на ключове за всички варианти ─────────────────────────────
all_keys = {}  # variant_num → list of correct letter indices
for v in range(1, NUM_VARIANTS + 1):
    variant = build_variant(v * 1000)
    all_keys[v] = [q[3] for q in variant]

# ── Документ ──────────────────────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(1.5)

# ── Заглавна страница ─────────────────────────────────────────────────────
h = doc.add_heading('Тест — Файлови системи и Shell скриптиране', 0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in h.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub = doc.add_paragraph('Операционни системи  ·  11А клас  ·  26 варианта')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in sub.runs:
    r.font.bold = True; r.font.size = Pt(13)

doc.add_paragraph()

# ── Скала за оценяване ────────────────────────────────────────────────────
sh = doc.add_heading('Скала за оценяване  (30 въпроса × 1 т.)', level=1)
for r in sh.runs:
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

scale_tbl = doc.add_table(rows=2, cols=6)
scale_tbl.style = 'Table Grid'
scale_headers = ['Верни отговори', '0 – 12', '13 – 16', '17 – 20', '21 – 25', '26 – 30']
scale_grades  = ['Оценка',         '2 (Слаб)', '3 (Среден)', '4 (Добър)', '5 (Мн. добър)', '6 (Отличен)']
bg_colors     = ['1F497D', 'FFCCCC', 'FFE5CC', 'FFFFCC', 'CCFFCC', 'CCECFF']
fg_colors     = [RGBColor(0xFF,0xFF,0xFF)] + [RGBColor(0x11,0x11,0x11)]*5

for ci, (h_text, g_text, bg, fg) in enumerate(zip(scale_headers, scale_grades, bg_colors, fg_colors)):
    hc = scale_tbl.rows[0].cells[ci]
    gc = scale_tbl.rows[1].cells[ci]
    hc.text = h_text
    gc.text = g_text
    for cell in (hc, gc):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg)
        tcPr.append(shd)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = fg

doc.add_paragraph()

# Обобщена таблица с ключовете на всички варианти
kh = doc.add_heading('Обобщен ключ — всички 26 варианта', level=1)
for r in kh.runs:
    r.font.color.rgb = RGBColor(0xC6, 0x20, 0x28)

info = doc.add_paragraph('Таблицата показва верните отговори (А/Б/В/Г) за всеки вариант и въпрос. САМО ЗА УЧИТЕЛЯ.')
for r in info.runs:
    r.font.italic = True; r.font.size = Pt(10)

doc.add_paragraph()

# Таблица: редове = варианти, колони = въпроси 1-30
tbl = doc.add_table(rows=NUM_VARIANTS + 1, cols=32)
tbl.style = 'Table Grid'

# Заглавен ред
hdr = tbl.rows[0].cells
hdr[0].text = 'Вар.'
hdr[1].text = 'Раздел'
set_cell_bg(hdr[0], '1F497D')
set_cell_bg(hdr[1], '1F497D')
for c in hdr[0].paragraphs[0].runs:
    c.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); c.font.bold = True; c.font.size = Pt(8)
for c in hdr[1].paragraphs[0].runs:
    c.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); c.font.bold = True; c.font.size = Pt(8)

for qi in range(30):
    hdr[qi+2].text = str(qi+1)
    set_cell_bg(hdr[qi+2], '1F497D')
    for r in hdr[qi+2].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.bold = True; r.font.size = Pt(8)

# Данни
for v in range(1, NUM_VARIANTS + 1):
    row = tbl.rows[v].cells
    row[0].text = str(v)
    row[1].text = '6+7'
    set_cell_bg(row[0], 'D9E1F2')
    set_cell_bg(row[1], 'D9E1F2')
    for c in row[0].paragraphs[0].runs:
        c.font.bold = True; c.font.size = Pt(8)
    for c in row[1].paragraphs[0].runs:
        c.font.size = Pt(8)
    for qi, letter_idx in enumerate(all_keys[v]):
        row[qi+2].text = LETTERS[letter_idx]
        bg = 'E2EFDA' if (qi % 2 == 0) else 'FFFFFF'
        set_cell_bg(row[qi+2], bg)
        for r in row[qi+2].paragraphs[0].runs:
            r.font.bold = True; r.font.size = Pt(8)

doc.add_page_break()

# ── Генериране на всички 26 варианта ─────────────────────────────────────
for v in range(1, NUM_VARIANTS + 1):
    variant = build_variant(v * 1000)

    # Заглавие на варианта
    doc.add_paragraph()
    vh = doc.add_heading(f'Вариант  {v}', level=1)
    vh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in vh.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        r.font.size = Pt(16)

    # Мета информация
    meta = doc.add_table(rows=1, cols=3)
    meta.style = 'Table Grid'
    meta.rows[0].cells[0].text = 'Операционни системи  ·  11А клас'
    meta.rows[0].cells[1].text = 'Три имена: _______________________________'
    meta.rows[0].cells[2].text = 'Оценка: _______'
    for cell in meta.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Скала за оценяване под метаданните
    sc = doc.add_paragraph()
    sc.paragraph_format.space_before = Pt(2)
    sc.paragraph_format.space_after  = Pt(4)
    sr = sc.add_run('Скала: 0–12 → 2  |  13–16 → 3  |  17–20 → 4  |  21–25 → 5  |  26–30 → 6')
    sr.font.size = Pt(8.5)
    sr.font.italic = True
    sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    sc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Въпроси
    current_section = ''
    for idx, (section, q, opts, correct) in enumerate(variant, 1):
        if section != current_section:
            current_section = section
            sh = doc.add_paragraph()
            run = sh.add_run(f'— {section} —')
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            sh.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Въпрос
        qp = doc.add_paragraph()
        qp.paragraph_format.space_before = Pt(3)
        qp.paragraph_format.space_after  = Pt(1)
        qr = qp.add_run(f'{idx}.  {q}')
        qr.font.bold = True
        qr.font.size = Pt(10)

        # Отговори — 2 колони за компактност
        opt_tbl = doc.add_table(rows=2, cols=2)
        for oi, opt_text in enumerate(opts):
            row_i = oi // 2
            col_i = oi % 2
            cell = opt_tbl.rows[row_i].cells[col_i]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            r = p.add_run(f'{LETTERS[oi]})  {opt_text}')
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Нова страница след всеки вариант (освен последния)
    if v < NUM_VARIANTS:
        doc.add_page_break()

# ── Запис ─────────────────────────────────────────────────────────────────
out = Path(r'C:\Users\Neli Nqgolova\Documents\Education\11а клас\ОС\Тест_26_варианта_ОС.docx')
doc.save(str(out))
print(f'✅ Записано: {out.name}')
print(f'   Варианти: {NUM_VARIANTS}')
print(f'   Въпроси на вариант: 30')
print(f'   Всеки вариант: разместени въпроси И разместени отговори')
